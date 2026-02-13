"""File parsing module for extracting text from various file formats."""
import io
from typing import Dict, Optional
import PyPDF2
import pdfplumber
import pandas as pd
from docx import Document
from utils.helpers import clean_text, sanitize_filename


class FileParser:
    """Handles parsing of different file formats."""
    
    @staticmethod
    def parse_pdf(file) -> str:
        """
        Parse PDF file and extract text.
        
        Args:
            file: Uploaded file object
        
        Returns:
            Extracted text content
        """
        text = ""
        
        try:
            # Try pdfplumber first (better for complex PDFs)
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except:
            # Fallback to PyPDF2
            try:
                file.seek(0)  # Reset file pointer
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            except Exception as e:
                raise Exception(f"Error parsing PDF: {str(e)}")
        
        return clean_text(text)
    
    @staticmethod
    def parse_excel(file) -> str:
        """
        Parse Excel file and extract text from all sheets.
        
        Args:
            file: Uploaded file object
        
        Returns:
            Extracted text content
        """
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file)
            text = ""
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                text += f"\n=== Sheet: {sheet_name} ===\n"
                text += df.to_string(index=False) + "\n"
            
            return clean_text(text)
        except Exception as e:
            raise Exception(f"Error parsing Excel file: {str(e)}")
    
    @staticmethod
    def parse_csv(file) -> str:
        """
        Parse CSV file and extract text.
        
        Args:
            file: Uploaded file object
        
        Returns:
            Extracted text content
        """
        try:
            df = pd.read_csv(file)
            text = df.to_string(index=False)
            return clean_text(text)
        except Exception as e:
            raise Exception(f"Error parsing CSV file: {str(e)}")
    
    @staticmethod
    def parse_docx(file) -> str:
        """
        Parse Word document and extract text.
        
        Args:
            file: Uploaded file object
        
        Returns:
            Extracted text content
        """
        try:
            doc = Document(file)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return clean_text(text)
        except Exception as e:
            raise Exception(f"Error parsing Word document: {str(e)}")
    
    @staticmethod
    def parse_txt(file) -> str:
        """
        Parse plain text file.
        
        Args:
            file: Uploaded file object
        
        Returns:
            Extracted text content
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    file.seek(0)
                    text = file.read().decode(encoding)
                    return clean_text(text)
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file with supported encodings")
        except Exception as e:
            raise Exception(f"Error parsing text file: {str(e)}")
    
    @staticmethod
    def parse_file(file, filename: str) -> Dict[str, str]:
        """
        Parse file based on its extension.
        
        Args:
            file: Uploaded file object
            filename: Original filename
        
        Returns:
            Dict with 'content' (extracted text) and 'file_type'
        """
        filename = sanitize_filename(filename)
        file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        try:
            if file_ext == 'pdf':
                content = FileParser.parse_pdf(file)
                file_type = 'pdf'
            elif file_ext in ['xlsx', 'xls']:
                content = FileParser.parse_excel(file)
                file_type = 'excel'
            elif file_ext == 'csv':
                content = FileParser.parse_csv(file)
                file_type = 'csv'
            elif file_ext in ['docx', 'doc']:
                content = FileParser.parse_docx(file)
                file_type = 'docx'
            elif file_ext == 'txt':
                content = FileParser.parse_txt(file)
                file_type = 'text'
            else:
                # Try as text file
                try:
                    content = FileParser.parse_txt(file)
                    file_type = 'text'
                except:
                    raise Exception(f"Unsupported file type: {file_ext}")
            
            if not content or len(content.strip()) == 0:
                raise Exception("No text content could be extracted from the file")
            
            return {
                'content': content,
                'file_type': file_type,
                'filename': filename
            }
        
        except Exception as e:
            raise Exception(f"Error parsing file '{filename}': {str(e)}")
