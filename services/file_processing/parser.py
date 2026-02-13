"""Strategy pattern for file parsing."""
import abc
import pdfplumber
import pandas as pd
from docx import Document
from typing import Protocol, Dict, Type
from domain.exceptions import ValidationError


class FileParsingStrategy(Protocol):
    """Protocol defining the file parsing strategy interface."""
    def parse(self, file_content: bytes, filename: str) -> str:
        """Parse file content and return text."""
        ...


class PdfParser:
    """Strategy for parsing PDF files."""
    def parse(self, file_content: bytes, filename: str) -> str:
        text = ""
        try:
            with pdfplumber.open(file_content) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            raise ValidationError(f"Failed to parse PDF: {str(e)}")
        return text


class ExcelParser:
    """Strategy for parsing Excel files."""
    def parse(self, file_content: bytes, filename: str) -> str:
        try:
            excel_file = pd.ExcelFile(file_content)
            text = ""
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                text += f"\n=== Sheet: {sheet_name} ===\n"
                text += df.to_string(index=False) + "\n"
            return text
        except Exception as e:
            raise ValidationError(f"Failed to parse Excel: {str(e)}")


class DocxParser:
    """Strategy for parsing Word documents."""
    def parse(self, file_content: bytes, filename: str) -> str:
        try:
            doc = Document(file_content)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text
        except Exception as e:
            raise ValidationError(f"Failed to parse Word doc: {str(e)}")


class TextParser:
    """Strategy for parsing plain text files."""
    def parse(self, file_content: bytes, filename: str) -> str:
        try:
            # Try decoding with common encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    # If it's a file-like object, read it
                    if hasattr(file_content, 'read'):
                        file_content.seek(0)
                        return file_content.read().decode(encoding)
                    # If it's bytes
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValidationError("Could not decode text file")
        except Exception as e:
            raise ValidationError(f"Failed to parse text file: {str(e)}")


class FileParserService:
    """Context for file parsing strategies."""
    
    def __init__(self):
        self._strategies: Dict[str, Type[FileParsingStrategy]] = {
            'pdf': PdfParser,
            'xlsx': ExcelParser,
            'xls': ExcelParser,
            'docx': DocxParser,
            'doc': DocxParser,
            'txt': TextParser,
            'csv': TextParser
        }

    def parse_file(self, file_obj, filename: str) -> str:
        """Parse a file using the appropriate strategy."""
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        strategy_class = self._strategies.get(ext, TextParser)
        strategy = strategy_class()
        
        # Handle Streamlit UploadedFile -> bytes/file-like
        return strategy.parse(file_obj, filename)
